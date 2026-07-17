#!/usr/bin/env python3
"""Generate all Day 026 notebooks: exercises 1-5, project, solution."""
import json
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DAY_DIR = ROOT / "02_automation" / "day_026"

_cid = 0


def cid():
    global _cid
    _cid += 1
    return f"c{_cid:04d}"


def md(source: str) -> dict:
    return {"cell_type": "markdown", "id": cid(), "metadata": {}, "source": source}


def code(source: str) -> dict:
    return {
        "cell_type": "code",
        "id": cid(),
        "metadata": {},
        "outputs": [],
        "execution_count": None,
        "source": source,
    }


def nb(cells: list) -> dict:
    return {
        "nbformat": 4,
        "nbformat_minor": 5,
        "metadata": {
            "kernelspec": {
                "display_name": "ai-course",
                "language": "python",
                "name": "ai-course",
            },
            "language_info": {"name": "python", "version": "3.11.0"},
        },
        "cells": cells,
    }


def write_nb(path: Path, cells: list):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(nb(cells), indent=1), encoding="utf-8")
    print(f"  wrote {path.relative_to(ROOT)}")


# ---------------------------------------------------------------------------
# Implementations (as they appear in generated notebooks)
# ---------------------------------------------------------------------------

READ_PDF_IMPL = """\
def read_pdf_pages(pdf_path: str) -> list[str]:
    reader = PdfReader(pdf_path)
    return [page.extract_text() or "" for page in reader.pages]"""

READ_DOCX_IMPL = """\
def read_docx_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\\n".join(p.text for p in doc.paragraphs)"""

CREATE_PDF_IMPL = """\
def create_pdf_report(title: str, sections: list[dict], output_path: str) -> None:
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(escape(title), styles["Title"]),
        Spacer(1, 0.25 * inch),
    ]
    for section in sections:
        story.append(Paragraph(escape(section["heading"]), styles["Heading1"]))
        story.append(Paragraph(escape(section["body"]), styles["Normal"]))
        story.append(Spacer(1, 0.15 * inch))
    doc.build(story)"""

CREATE_DOCX_IMPL = """\
def create_docx_report(title: str, sections: list[dict], output_path: str) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for section in sections:
        doc.add_heading(section["heading"], level=1)
        doc.add_paragraph(section["body"])
    doc.save(output_path)"""

AI_GEN_IMPL = """\
def ai_generate_section(
    topic: str,
    data_snippet: str,
    model: str = "llama3.2",
) -> dict:
    response = ollama.chat(
        model=model,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a professional report writer. "
                    "Generate a concise report section from the data provided. "
                    'Return JSON with exactly two keys: "heading" (a short title string) '
                    'and "body" (2-4 sentences of professional analysis). '
                    "Return only valid JSON, no explanation."
                ),
            },
            {
                "role": "user",
                "content": (
                    f"Topic: {topic}\\n\\n"
                    f"Data:\\n{data_snippet[:500]}\\n\\n"
                    "Generate a report section:"
                ),
            },
        ],
        format="json",
    )
    raw = response["message"]["content"]
    try:
        result = json.loads(raw)
        return {
            "heading": str(result.get("heading", topic)),
            "body": str(result.get("body", "")),
        }
    except Exception:
        return {"heading": topic, "body": raw}"""

AUTO_REPORTER_IMPL = """\
class AutoReporter:
    def __init__(self, model: str = "llama3.2"):
        self.model = model

    def generate_sections(self, data: dict) -> list[dict]:
        sections = []
        for key, value in data.items():
            snippet = (
                json.dumps(value, indent=2)
                if isinstance(value, (dict, list))
                else str(value)
            )
            section = ai_generate_section(key, snippet, model=self.model)
            sections.append(section)
        return sections

    def to_pdf(self, title: str, sections: list[dict], output_path: str) -> None:
        create_pdf_report(title, sections, output_path)

    def to_docx(self, title: str, sections: list[dict], output_path: str) -> None:
        create_docx_report(title, sections, output_path)

    def generate_report(self, data: dict, title: str, output_dir: str = ".") -> dict:
        import os
        sections = self.generate_sections(data)
        pdf_path = os.path.join(output_dir, "report.pdf")
        docx_path = os.path.join(output_dir, "report.docx")
        self.to_pdf(title, sections, pdf_path)
        self.to_docx(title, sections, docx_path)
        return {"pdf": pdf_path, "docx": docx_path, "sections": sections}\
"""

ALL_IMPLS = "\n\n\n".join([
    READ_PDF_IMPL, READ_DOCX_IMPL, CREATE_PDF_IMPL, CREATE_DOCX_IMPL, AI_GEN_IMPL,
])


# ---------------------------------------------------------------------------
# Exercise 01 — read_pdf_pages
# ---------------------------------------------------------------------------

def ex01():
    global _cid; _cid = 0
    return [
        md(
            "# Day 026 — Exercise 1: read_pdf_pages\n\n"
            "**What you'll build:** `read_pdf_pages(pdf_path)` — reads a PDF with `pypdf` "
            "and returns a list of strings, one per page.\n\n"
            "**Why it matters:** Before AI can analyse a PDF (contract, invoice, report), "
            "you need the text. `pypdf` is the go-to stdlib-free option for text-layer PDFs. "
            "`page.extract_text() or ''` guards against pages that return None "
            "(images-only pages or pages with no text layer)."
        ),
        code("from pypdf import PdfReader"),
        md("## Your Implementation"),
        code(
            "def read_pdf_pages(pdf_path: str) -> list[str]:\n"
            '    """\n'
            "    Extract text from each page of a PDF file.\n\n"
            "    Args:\n"
            "        pdf_path: Path to the PDF file.\n\n"
            "    Returns:\n"
            "        List of strings, one per page. Empty string for pages with no text.\n"
            '    """\n'
            "    # TODO: reader = PdfReader(pdf_path)\n"
            "    # TODO: return [page.extract_text() or \"\" for page in reader.pages]\n"
            "    pass"
        ),
        md("## Check Your Work"),
        code(
            "from reportlab.pdfgen import canvas as _rl_canvas\n"
            "from reportlab.lib.pagesizes import letter as _letter\n"
            "import tempfile, os\n"
            "\n"
            "\n"
            "def _run_checks():\n"
            "    total = 5\n"
            "    passed = 0\n"
            "\n"
            "    # Create a 2-page test PDF\n"
            "    tmp = tempfile.mktemp(suffix='.pdf')\n"
            "    _c = _rl_canvas.Canvas(tmp, pagesize=_letter)\n"
            "    _c.setFont('Helvetica', 12)\n"
            "    _c.drawString(72, 750, 'Test Page One Content')\n"
            "    _c.showPage()\n"
            "    _c.drawString(72, 750, 'Test Page Two Content')\n"
            "    _c.save()\n"
            "\n"
            "    pages = None\n"
            "\n"
            "    # Check 1: defined\n"
            "    try:\n"
            "        assert 'read_pdf_pages' in globals()\n"
            "        passed += 1; print('\\u2705 Check 1: read_pdf_pages defined')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 1: {e}')\n"
            "        os.unlink(tmp)\n"
            "        print(f'\\nScore: {passed}/{total}')\n"
            "        return\n"
            "\n"
            "    # Check 2: returns a list\n"
            "    try:\n"
            "        pages = read_pdf_pages(tmp)\n"
            "        assert isinstance(pages, list), f'expected list, got {type(pages)}'\n"
            "        passed += 1; print('\\u2705 Check 2: returns a list')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 2: {e}')\n"
            "\n"
            "    # Check 3: list has 2 items (2-page PDF)\n"
            "    try:\n"
            "        assert pages is not None, 'pages is None (Check 2 failed)'\n"
            "        assert len(pages) == 2, f'expected 2 pages, got {len(pages)}'\n"
            "        passed += 1; print(f'\\u2705 Check 3: {len(pages)} pages extracted')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 3: {e}')\n"
            "\n"
            "    # Check 4: all items are strings\n"
            "    try:\n"
            "        assert pages is not None, 'pages is None'\n"
            "        bad = [p for p in pages if not isinstance(p, str)]\n"
            "        assert not bad, f'non-string pages: {bad}'\n"
            "        passed += 1; print('\\u2705 Check 4: all page items are strings')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 4: {e}')\n"
            "\n"
            "    # Check 5: at least one page contains expected text\n"
            "    try:\n"
            "        assert pages is not None, 'pages is None'\n"
            "        combined = ' '.join(pages)\n"
            "        assert 'Test Page' in combined, \\\n"
            "            f\"'Test Page' not found in extracted text: {combined!r}\"\n"
            "        passed += 1; print('\\u2705 Check 5: extracted text contains expected content')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 5: {e}')\n"
            "\n"
            "    try:\n"
            "        os.unlink(tmp)\n"
            "    except Exception:\n"
            "        pass\n"
            "\n"
            "    if passed == total:\n"
            "        print('\\U0001f389 Exercise complete!')\n"
            "    print(f'\\nScore: {passed}/{total}')\n"
            "\n"
            "\n"
            "_run_checks()"
        ),
        md(
            "## Solution\n\n"
            "<details>\n"
            "<summary>Click to reveal</summary>\n\n"
            "```python\n"
            + READ_PDF_IMPL + "\n"
            "```\n\n"
            "</details>"
        ),
    ]


# ---------------------------------------------------------------------------
# Exercise 02 — read_docx_text
# ---------------------------------------------------------------------------

def ex02():
    global _cid; _cid = 100
    return [
        md(
            "# Day 026 — Exercise 2: read_docx_text\n\n"
            "**What you'll build:** `read_docx_text(docx_path)` — opens a `.docx` file with "
            "`python-docx` and returns all paragraph text joined with newlines.\n\n"
            "**Why it matters:** Word documents are common in business workflows. "
            "`python-docx` exposes every paragraph as an object with a `.text` attribute. "
            "Joining with `\\n` gives clean line-separated text that an LLM can read directly."
        ),
        code("from docx import Document"),
        md("## Your Implementation"),
        code(
            "def read_docx_text(docx_path: str) -> str:\n"
            '    """\n'
            "    Extract all paragraph text from a DOCX file.\n\n"
            "    Args:\n"
            "        docx_path: Path to the .docx file.\n\n"
            "    Returns:\n"
            "        All paragraph text joined with newlines (includes headings and body).\n"
            '    """\n'
            "    # TODO: doc = Document(docx_path)\n"
            '    # TODO: return "\\n".join(p.text for p in doc.paragraphs)\n'
            "    pass"
        ),
        md("## Check Your Work"),
        code(
            "from docx import Document as _Document\n"
            "import tempfile, os\n"
            "\n"
            "\n"
            "def _run_checks():\n"
            "    total = 5\n"
            "    passed = 0\n"
            "\n"
            "    # Create a test DOCX\n"
            "    tmp = tempfile.mktemp(suffix='.docx')\n"
            "    _doc = _Document()\n"
            "    _doc.add_heading('Meeting Summary', level=0)\n"
            "    _doc.add_paragraph('The team discussed the quarterly results.')\n"
            "    _doc.add_paragraph('Action items were assigned to each member.')\n"
            "    _doc.save(tmp)\n"
            "\n"
            "    text = None\n"
            "\n"
            "    # Check 1: defined\n"
            "    try:\n"
            "        assert 'read_docx_text' in globals()\n"
            "        passed += 1; print('\\u2705 Check 1: read_docx_text defined')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 1: {e}')\n"
            "        os.unlink(tmp)\n"
            "        print(f'\\nScore: {passed}/{total}')\n"
            "        return\n"
            "\n"
            "    # Check 2: returns a string\n"
            "    try:\n"
            "        text = read_docx_text(tmp)\n"
            "        assert isinstance(text, str), f'expected str, got {type(text)}'\n"
            "        passed += 1; print('\\u2705 Check 2: returns a string')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 2: {e}')\n"
            "\n"
            "    # Check 3: contains the heading text\n"
            "    try:\n"
            "        assert text is not None, 'text is None (Check 2 failed)'\n"
            "        assert 'Meeting Summary' in text, \\\n"
            "            f\"'Meeting Summary' not in text: {text!r}\"\n"
            "        passed += 1; print(\"\\u2705 Check 3: contains heading 'Meeting Summary'\")\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 3: {e}')\n"
            "\n"
            "    # Check 4: contains paragraph body text\n"
            "    try:\n"
            "        assert text is not None, 'text is None'\n"
            "        assert 'quarterly results' in text, \\\n"
            "            f\"'quarterly results' not in text: {text!r}\"\n"
            "        passed += 1; print(\"\\u2705 Check 4: contains paragraph body text\")\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 4: {e}')\n"
            "\n"
            "    # Check 5: multiple paragraphs are separated (newline present)\n"
            "    try:\n"
            "        assert text is not None, 'text is None'\n"
            "        assert '\\n' in text, \\\n"
            "            f'expected newlines between paragraphs, got {text!r}'\n"
            "        passed += 1; print('\\u2705 Check 5: paragraphs separated by newlines')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 5: {e}')\n"
            "\n"
            "    try:\n"
            "        os.unlink(tmp)\n"
            "    except Exception:\n"
            "        pass\n"
            "\n"
            "    if passed == total:\n"
            "        print('\\U0001f389 Exercise complete!')\n"
            "    print(f'\\nScore: {passed}/{total}')\n"
            "\n"
            "\n"
            "_run_checks()"
        ),
        md(
            "## Solution\n\n"
            "<details>\n"
            "<summary>Click to reveal</summary>\n\n"
            "```python\n"
            + READ_DOCX_IMPL + "\n"
            "```\n\n"
            "</details>"
        ),
    ]


# ---------------------------------------------------------------------------
# Exercise 03 — create_pdf_report
# ---------------------------------------------------------------------------

def ex03():
    global _cid; _cid = 200
    return [
        md(
            "# Day 026 — Exercise 3: create_pdf_report\n\n"
            "**What you'll build:** `create_pdf_report(title, sections, output_path)` — "
            "generates a formatted PDF with a title heading and styled sections using "
            "`reportlab.platypus`.\n\n"
            "**Why it matters:** `reportlab.platypus` (Paragraph, SimpleDocTemplate) handles "
            "word-wrap, page breaks, and font styling automatically — you describe *what*, "
            "not *where*. The `escape()` call prevents XML parsing errors when the text "
            "contains `<`, `>`, or `&`."
        ),
        code(
            "from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer\n"
            "from reportlab.lib.styles import getSampleStyleSheet\n"
            "from reportlab.lib.pagesizes import letter\n"
            "from reportlab.lib.units import inch\n"
            "from xml.sax.saxutils import escape"
        ),
        md("## Your Implementation"),
        code(
            "def create_pdf_report(\n"
            "    title: str,\n"
            "    sections: list[dict],\n"
            "    output_path: str,\n"
            ") -> None:\n"
            '    """\n'
            "    Create a formatted PDF report using reportlab platypus.\n\n"
            "    Args:\n"
            "        title:       Report title (rendered as Title style).\n"
            "        sections:    List of {'heading': str, 'body': str} dicts.\n"
            "        output_path: File path to write the .pdf.\n\n"
            "    Note: escape() prevents XML parse errors on special chars in text.\n"
            '    """\n'
            "    # TODO: doc = SimpleDocTemplate(output_path, pagesize=letter)\n"
            "    # TODO: styles = getSampleStyleSheet()\n"
            "    # TODO: story = [\n"
            "    #           Paragraph(escape(title), styles['Title']),\n"
            "    #           Spacer(1, 0.25 * inch),\n"
            "    #       ]\n"
            "    # TODO: for section in sections:\n"
            "    #           story.append(Paragraph(escape(section['heading']), styles['Heading1']))\n"
            "    #           story.append(Paragraph(escape(section['body']), styles['Normal']))\n"
            "    #           story.append(Spacer(1, 0.15 * inch))\n"
            "    # TODO: doc.build(story)\n"
            "    pass"
        ),
        md("## Check Your Work"),
        code(
            "from pypdf import PdfReader as _PdfReader\n"
            "import tempfile, os\n"
            "\n"
            "\n"
            "def _run_checks():\n"
            "    total = 5\n"
            "    passed = 0\n"
            "\n"
            "    SECTIONS = [\n"
            "        {'heading': 'Introduction', 'body': 'This report covers Q1 results.'},\n"
            "        {'heading': 'Findings',     'body': 'Revenue increased by 15 percent year over year.'},\n"
            "    ]\n"
            "    tmp = tempfile.mktemp(suffix='.pdf')\n"
            "\n"
            "    # Check 1: defined\n"
            "    try:\n"
            "        assert 'create_pdf_report' in globals()\n"
            "        passed += 1; print('\\u2705 Check 1: create_pdf_report defined')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 1: {e}')\n"
            "        print(f'\\nScore: {passed}/{total}')\n"
            "        return\n"
            "\n"
            "    # Check 2: creates file without error\n"
            "    try:\n"
            "        create_pdf_report('Quarterly Report', SECTIONS, tmp)\n"
            "        assert os.path.exists(tmp), f'file not created at {tmp}'\n"
            "        passed += 1; print(f'\\u2705 Check 2: file created at {tmp}')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 2: {e}')\n"
            "\n"
            "    # Check 3: file starts with PDF magic bytes\n"
            "    try:\n"
            "        assert os.path.exists(tmp), 'file does not exist (Check 2 failed)'\n"
            "        with open(tmp, 'rb') as f:\n"
            "            header = f.read(5)\n"
            "        assert header == b'%PDF-', f'not a PDF: {header!r}'\n"
            "        passed += 1; print(f'\\u2705 Check 3: file has PDF magic bytes')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 3: {e}')\n"
            "\n"
            "    # Check 4: file is non-trivial size (> 500 bytes)\n"
            "    try:\n"
            "        assert os.path.exists(tmp), 'file does not exist'\n"
            "        size = os.path.getsize(tmp)\n"
            "        assert size > 500, f'PDF too small ({size} bytes) — probably empty'\n"
            "        passed += 1; print(f'\\u2705 Check 4: file is {size:,} bytes')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 4: {e}')\n"
            "\n"
            "    # Check 5: pypdf can parse it (valid structure)\n"
            "    try:\n"
            "        assert os.path.exists(tmp), 'file does not exist'\n"
            "        reader = _PdfReader(tmp)\n"
            "        n = len(reader.pages)\n"
            "        assert n >= 1, f'expected at least 1 page, got {n}'\n"
            "        passed += 1; print(f'\\u2705 Check 5: valid PDF with {n} page(s)')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 5: {e}')\n"
            "\n"
            "    try:\n"
            "        os.unlink(tmp)\n"
            "    except Exception:\n"
            "        pass\n"
            "\n"
            "    if passed == total:\n"
            "        print('\\U0001f389 Exercise complete!')\n"
            "    print(f'\\nScore: {passed}/{total}')\n"
            "\n"
            "\n"
            "_run_checks()"
        ),
        md(
            "## Solution\n\n"
            "<details>\n"
            "<summary>Click to reveal</summary>\n\n"
            "```python\n"
            + CREATE_PDF_IMPL + "\n"
            "```\n\n"
            "</details>"
        ),
    ]


# ---------------------------------------------------------------------------
# Exercise 04 — create_docx_report
# ---------------------------------------------------------------------------

def ex04():
    global _cid; _cid = 300
    return [
        md(
            "# Day 026 — Exercise 4: create_docx_report\n\n"
            "**What you'll build:** `create_docx_report(title, sections, output_path)` — "
            "creates a Word document with a title and styled section headings using "
            "`python-docx`.\n\n"
            "**Why it matters:** DOCX is editable after generation — clients can open, "
            "annotate, and reformat in Word. `doc.add_heading(text, level=0)` gives the "
            "document title; `level=1` gives section headings. "
            "`doc.add_paragraph(body)` adds the body text."
        ),
        code("from docx import Document"),
        md("## Your Implementation"),
        code(
            "def create_docx_report(\n"
            "    title: str,\n"
            "    sections: list[dict],\n"
            "    output_path: str,\n"
            ") -> None:\n"
            '    """\n'
            "    Create a DOCX report with a title and styled sections.\n\n"
            "    Args:\n"
            "        title:       Report title (level-0 heading).\n"
            "        sections:    List of {'heading': str, 'body': str} dicts.\n"
            "        output_path: File path to write the .docx.\n"
            '    """\n'
            "    # TODO: doc = Document()\n"
            "    # TODO: doc.add_heading(title, level=0)\n"
            "    # TODO: for section in sections:\n"
            "    #           doc.add_heading(section['heading'], level=1)\n"
            "    #           doc.add_paragraph(section['body'])\n"
            "    # TODO: doc.save(output_path)\n"
            "    pass"
        ),
        md("## Check Your Work"),
        code(
            "from docx import Document as _Document\n"
            "import tempfile, os\n"
            "\n"
            "\n"
            "def _run_checks():\n"
            "    total = 5\n"
            "    passed = 0\n"
            "\n"
            "    SECTIONS = [\n"
            "        {'heading': 'Q1 Summary', 'body': 'Sales rose 12 percent in the first quarter.'},\n"
            "        {'heading': 'Outlook',    'body': 'We expect continued growth in Q2.'},\n"
            "    ]\n"
            "    tmp = tempfile.mktemp(suffix='.docx')\n"
            "\n"
            "    # Check 1: defined\n"
            "    try:\n"
            "        assert 'create_docx_report' in globals()\n"
            "        passed += 1; print('\\u2705 Check 1: create_docx_report defined')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 1: {e}')\n"
            "        print(f'\\nScore: {passed}/{total}')\n"
            "        return\n"
            "\n"
            "    # Check 2: creates file without error\n"
            "    try:\n"
            "        create_docx_report('Annual Report', SECTIONS, tmp)\n"
            "        assert os.path.exists(tmp), f'file not created at {tmp}'\n"
            "        passed += 1; print(f'\\u2705 Check 2: file created at {tmp}')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 2: {e}')\n"
            "\n"
            "    doc = None\n"
            "\n"
            "    # Check 3: python-docx can open the file\n"
            "    try:\n"
            "        assert os.path.exists(tmp), 'file does not exist (Check 2 failed)'\n"
            "        doc = _Document(tmp)\n"
            "        assert doc is not None\n"
            "        passed += 1; print('\\u2705 Check 3: file is a valid DOCX')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 3: {e}')\n"
            "\n"
            "    # Check 4: title appears in paragraphs\n"
            "    try:\n"
            "        assert doc is not None, 'doc is None (Check 3 failed)'\n"
            "        texts = [p.text for p in doc.paragraphs]\n"
            "        assert any('Annual Report' in t for t in texts), \\\n"
            "            f\"'Annual Report' not found in paragraphs: {texts}\"\n"
            "        passed += 1; print(\"\\u2705 Check 4: title 'Annual Report' in document\")\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 4: {e}')\n"
            "\n"
            "    # Check 5: section heading appears in paragraphs\n"
            "    try:\n"
            "        assert doc is not None, 'doc is None'\n"
            "        texts = [p.text for p in doc.paragraphs]\n"
            "        assert any('Q1 Summary' in t for t in texts), \\\n"
            "            f\"'Q1 Summary' not found in paragraphs: {texts}\"\n"
            "        passed += 1; print(\"\\u2705 Check 5: section heading 'Q1 Summary' in document\")\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 5: {e}')\n"
            "\n"
            "    try:\n"
            "        os.unlink(tmp)\n"
            "    except Exception:\n"
            "        pass\n"
            "\n"
            "    if passed == total:\n"
            "        print('\\U0001f389 Exercise complete!')\n"
            "    print(f'\\nScore: {passed}/{total}')\n"
            "\n"
            "\n"
            "_run_checks()"
        ),
        md(
            "## Solution\n\n"
            "<details>\n"
            "<summary>Click to reveal</summary>\n\n"
            "```python\n"
            + CREATE_DOCX_IMPL + "\n"
            "```\n\n"
            "</details>"
        ),
    ]


# ---------------------------------------------------------------------------
# Exercise 05 — ai_generate_section
# ---------------------------------------------------------------------------

def ex05():
    global _cid; _cid = 400
    return [
        md(
            "# Day 026 — Exercise 5: ai_generate_section\n\n"
            "**What you'll build:** `ai_generate_section(topic, data_snippet, model)` — "
            "calls the LLM with `format='json'` to generate a `{heading, body}` dict "
            "from raw data text.\n\n"
            "**Why it matters:** This is the bridge between raw data and a human-readable "
            "report section. The schema-guided approach (Day 4) gives you structured output "
            "you can pass directly to `create_pdf_report` or `create_docx_report`. "
            "The `json.loads` + fallback pattern ensures you always get a usable dict "
            "even if the LLM returns slightly malformed JSON."
        ),
        code("import ollama\nimport json"),
        md("## Your Implementation"),
        code(
            "def ai_generate_section(\n"
            "    topic: str,\n"
            "    data_snippet: str,\n"
            '    model: str = "llama3.2",\n'
            ") -> dict:\n"
            '    """\n'
            "    Generate a report section {heading, body} from raw data using the LLM.\n\n"
            "    Args:\n"
            "        topic:        Section topic (used as fallback heading if LLM fails).\n"
            "        data_snippet: Raw data text for the LLM to analyse.\n"
            "        model:        Ollama model name.\n\n"
            "    Returns:\n"
            "        dict with 'heading' (str) and 'body' (str). Never raises.\n"
            '    """\n'
            "    # TODO: ollama.chat(model=model, messages=[system, user], format='json')\n"
            "    # TODO: system: 'professional report writer; return JSON {heading, body}'\n"
            "    # TODO: user: f'Topic: {topic}\\n\\nData:\\n{data_snippet[:500]}'\n"
            "    # TODO: try: result = json.loads(response['message']['content'])\n"
            "    # TODO:      return {'heading': str(result.get('heading', topic)),\n"
            "    # TODO:              'body':    str(result.get('body', ''))}\n"
            "    # TODO: except Exception: return {'heading': topic, 'body': raw}\n"
            "    pass"
        ),
        md("## Check Your Work"),
        code(
            "def _run_checks():\n"
            "    total = 5\n"
            "    passed = 0\n"
            "\n"
            "    # Check 1: defined\n"
            "    try:\n"
            "        assert 'ai_generate_section' in globals()\n"
            "        passed += 1; print('\\u2705 Check 1: ai_generate_section defined')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 1: {e}')\n"
            "        print(f'\\nScore: {passed}/{total}')\n"
            "        return\n"
            "\n"
            "    result = None\n"
            "\n"
            "    # Check 2: returns a dict (1 LLM call)\n"
            "    try:\n"
            "        result = ai_generate_section(\n"
            "            'Sales Performance',\n"
            "            'Q1: 1.2M revenue, up 15%% YoY. Top product: Widget A.',\n"
            "        )\n"
            "        assert isinstance(result, dict), f'expected dict, got {type(result)}'\n"
            "        passed += 1; print('\\u2705 Check 2: returns a dict')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 2: {e}')\n"
            "\n"
            "    # Check 3: dict has 'heading' key with non-empty string\n"
            "    try:\n"
            "        assert result is not None, 'result is None (Check 2 failed)'\n"
            "        assert 'heading' in result, f\"missing 'heading' key: {result}\"\n"
            "        assert isinstance(result['heading'], str) and len(result['heading']) > 0, \\\n"
            "            f\"heading should be non-empty str: {result['heading']!r}\"\n"
            "        passed += 1; print(f\"\\u2705 Check 3: heading={result['heading']!r}\")\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 3: {e}')\n"
            "\n"
            "    # Check 4: dict has 'body' key with non-empty string\n"
            "    try:\n"
            "        assert result is not None, 'result is None'\n"
            "        assert 'body' in result, f\"missing 'body' key: {result}\"\n"
            "        assert isinstance(result['body'], str) and len(result['body']) > 0, \\\n"
            "            f\"body should be non-empty str: {result['body']!r}\"\n"
            "        passed += 1; print(f\"\\u2705 Check 4: body is {len(result['body'])} chars\")\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 4: {e}')\n"
            "\n"
            "    # Check 5: works on a different topic (1 LLM call)\n"
            "    try:\n"
            "        result2 = ai_generate_section(\n"
            "            'Market Overview',\n"
            "            'Total market: 50B. Our share: 2.4%%. Three main competitors.',\n"
            "        )\n"
            "        assert isinstance(result2, dict), f'expected dict, got {type(result2)}'\n"
            "        assert 'heading' in result2 and 'body' in result2\n"
            "        passed += 1; print('\\u2705 Check 5: works on different topic')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 5: {e}')\n"
            "\n"
            "    if passed == total:\n"
            "        print('\\U0001f389 Exercise complete!')\n"
            "    print(f'\\nScore: {passed}/{total}')\n"
            "\n"
            "\n"
            "_run_checks()"
        ),
        md(
            "## Solution\n\n"
            "<details>\n"
            "<summary>Click to reveal</summary>\n\n"
            "```python\n"
            + AI_GEN_IMPL + "\n"
            "```\n\n"
            "</details>"
        ),
    ]


# ---------------------------------------------------------------------------
# Project notebook (student template — NOT executed by gate)
# ---------------------------------------------------------------------------

def project_nb():
    global _cid; _cid = 500
    return [
        md(
            "# Day 026 Project: AutoReporter\n\n"
            "## What You're Building\n\n"
            "An `AutoReporter` class that takes a dict of raw data, uses the LLM to "
            "draft a section per key, and writes both a PDF and a DOCX report.\n\n"
            "Pipeline: `data dict → generate_sections (AI) → to_pdf + to_docx → files`\n\n"
            "## Project Requirements\n\n"
            "1. Implement `AutoReporter` with:\n"
            "   - `generate_sections(data: dict) -> list[dict]` — one AI call per key\n"
            "   - `to_pdf(title, sections, output_path)` — write PDF\n"
            "   - `to_docx(title, sections, output_path)` — write DOCX\n"
            "   - `generate_report(data, title, output_dir) -> dict` — full pipeline\n"
            "2. Run `reporter.generate_report(SAMPLE_DATA, 'Q1 Report', '/tmp')` and store "
            "as `result`\n"
            "3. Print the output file paths\n\n"
            "**Deliverable:** PDF and DOCX report files auto-generated from raw data."
        ),
        code(
            "import json, os\n"
            "from pypdf import PdfReader\n"
            "from docx import Document\n"
            "from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer\n"
            "from reportlab.lib.styles import getSampleStyleSheet\n"
            "from reportlab.lib.pagesizes import letter\n"
            "from reportlab.lib.units import inch\n"
            "from xml.sax.saxutils import escape\n"
            "import ollama"
        ),
        md("## Provided: All Helper Functions"),
        code(ALL_IMPLS),
        md(
            "## Your Implementation\n\n"
            "Implement `AutoReporter` by wiring the helper functions together."
        ),
        code(
            "class AutoReporter:\n"
            "    def __init__(self, model: str = 'llama3.2'):\n"
            "        self.model = model\n"
            "\n"
            "    def generate_sections(self, data: dict) -> list[dict]:\n"
            "        # TODO: for key, value in data.items():\n"
            "        #           snippet = json.dumps(value, indent=2) if isinstance(value, (dict, list)) else str(value)\n"
            "        #           section = ai_generate_section(key, snippet, model=self.model)\n"
            "        #           sections.append(section)\n"
            "        pass\n"
            "\n"
            "    def to_pdf(self, title: str, sections: list[dict], output_path: str) -> None:\n"
            "        # TODO: create_pdf_report(title, sections, output_path)\n"
            "        pass\n"
            "\n"
            "    def to_docx(self, title: str, sections: list[dict], output_path: str) -> None:\n"
            "        # TODO: create_docx_report(title, sections, output_path)\n"
            "        pass\n"
            "\n"
            "    def generate_report(self, data: dict, title: str, output_dir: str = '.') -> dict:\n"
            "        # TODO: sections = self.generate_sections(data)\n"
            "        # TODO: pdf_path = os.path.join(output_dir, 'report.pdf')\n"
            "        # TODO: docx_path = os.path.join(output_dir, 'report.docx')\n"
            "        # TODO: self.to_pdf(title, sections, pdf_path)\n"
            "        # TODO: self.to_docx(title, sections, docx_path)\n"
            "        # TODO: return {'pdf': pdf_path, 'docx': docx_path, 'sections': sections}\n"
            "        pass"
        ),
        md("## Sample Data"),
        code(
            "SAMPLE_DATA = {\n"
            "    'Sales Performance': 'Q1 revenue: 1.2M, up 15 percent YoY. Top product: Widget A with 38 percent share.',\n"
            "    'Market Overview':   'Total market: 50B. Our share: 2.4 percent. Three main competitors.',\n"
            "}\n"
        ),
        md("## Run the Reporter"),
        code(
            "# reporter = AutoReporter()\n"
            "# result = reporter.generate_report(SAMPLE_DATA, 'Q1 Business Report', '/tmp')\n"
            "# print(f\"PDF:  {result['pdf']}\")\n"
            "# print(f\"DOCX: {result['docx']}\")\n"
            "# for s in result['sections']:\n"
            "#     print(f\"  - {s['heading']}\")"
        ),
        md("## Checks"),
        code(
            "def _run_project_checks():\n"
            "    total = 5\n"
            "    passed = 0\n"
            "\n"
            "    # Check 1: AutoReporter defined with required methods\n"
            "    try:\n"
            "        assert 'AutoReporter' in globals()\n"
            "        for m in ('generate_sections', 'to_pdf', 'to_docx', 'generate_report'):\n"
            "            assert hasattr(AutoReporter, m), f'AutoReporter missing: {m}'\n"
            "        passed += 1; print('\\u2705 Check 1: AutoReporter has all required methods')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 1: {e}')\n"
            "\n"
            "    # Check 2: reporter is an AutoReporter instance\n"
            "    try:\n"
            "        assert 'reporter' in globals()\n"
            "        assert isinstance(reporter, AutoReporter)\n"
            "        passed += 1; print('\\u2705 Check 2: reporter is an AutoReporter')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 2: {e}')\n"
            "\n"
            "    # Check 3: result is a dict with required keys\n"
            "    try:\n"
            "        assert 'result' in globals()\n"
            "        for k in ('pdf', 'docx', 'sections'):\n"
            "            assert k in result, f\"result missing '{k}': {list(result)}\"\n"
            "        passed += 1; print('\\u2705 Check 3: result has pdf/docx/sections keys')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 3: {e}')\n"
            "\n"
            "    # Check 4: PDF file exists\n"
            "    try:\n"
            "        assert 'result' in globals()\n"
            "        pdf = result.get('pdf', '')\n"
            "        assert os.path.exists(pdf), f'PDF not found at {pdf}'\n"
            "        size = os.path.getsize(pdf)\n"
            "        assert size > 100, f'PDF too small ({size} bytes)'\n"
            "        passed += 1; print(f'\\u2705 Check 4: PDF exists ({size:,} bytes)')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 4: {e}')\n"
            "\n"
            "    # Check 5: DOCX exists and is readable\n"
            "    try:\n"
            "        assert 'result' in globals()\n"
            "        from docx import Document as _D\n"
            "        docx = result.get('docx', '')\n"
            "        assert os.path.exists(docx), f'DOCX not found at {docx}'\n"
            "        doc = _D(docx)\n"
            "        assert len(doc.paragraphs) > 0\n"
            "        passed += 1; print('\\u2705 Check 5: DOCX exists and is readable')\n"
            "    except Exception as e:\n"
            "        print(f'\\u274c Check 5: {e}')\n"
            "\n"
            "    if passed == total:\n"
            "        print('\\U0001f389 Project complete!')\n"
            "    print(f'\\nScore: {passed}/{total}')\n"
            "\n"
            "\n"
            "_run_project_checks()"
        ),
        md(
            "## Bonus Challenges\n\n"
            "- Add a `cover_page(title, subtitle, date)` helper that prepends a styled "
            "cover page to the PDF using reportlab canvas or platypus\n"
            "- Add a `summarize_report(sections, model) -> str` method that asks the LLM "
            "to write a one-paragraph executive summary across all sections\n"
            "- Extend `generate_sections` to accept a list instead of a dict, using each "
            "item's 'topic' key\n"
            "- Add an `append_table(doc, rows: list[list[str]])` helper using "
            "`python-docx`'s `add_table` API\n"
            "- Handle the case where `ai_generate_section` returns an empty body by retrying "
            "once with a more explicit prompt"
        ),
    ]


# ---------------------------------------------------------------------------
# Solution notebook (executed by gate)
# ---------------------------------------------------------------------------

def solution_nb():
    global _cid; _cid = 600

    imports = (
        "import json, os, tempfile\n"
        "from pypdf import PdfReader\n"
        "from docx import Document\n"
        "from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer\n"
        "from reportlab.lib.styles import getSampleStyleSheet\n"
        "from reportlab.lib.pagesizes import letter\n"
        "from reportlab.lib.units import inch\n"
        "from xml.sax.saxutils import escape\n"
        "import ollama"
    )

    all_fns = imports + "\n\n\n" + ALL_IMPLS + "\n\n\n" + AUTO_REPORTER_IMPL

    sample_data = (
        "SAMPLE_DATA = {\n"
        "    'Sales Performance': 'Q1 revenue: 1.2M, up 15 percent YoY. Top product: Widget A with 38 percent share.',\n"
        "    'Market Overview':   'Total market: 50B. Our share: 2.4 percent. Three main competitors.',\n"
        "}"
    )

    return [
        md(
            "# Day 026 Project Solution — AutoReporter\n\n"
            "An `AutoReporter` that turns raw data dicts into styled PDF and DOCX reports "
            "using AI-generated section content."
        ),
        code(all_fns),
        code(sample_data),
        md("## Action 1 — Generate Report Sections"),
        code(
            "reporter = AutoReporter()\n"
            "print('Generating report sections (2 AI calls)...')\n"
            "sections = reporter.generate_sections(SAMPLE_DATA)\n"
            "print(f'Generated {len(sections)} section(s):')\n"
            "for s in sections:\n"
            "    print(f\"  - {s['heading']}\")"
        ),
        md("## Action 2 — Write PDF and DOCX"),
        code(
            "tmp_dir = tempfile.gettempdir()\n"
            "pdf_path  = os.path.join(tmp_dir, 'day026_report.pdf')\n"
            "docx_path = os.path.join(tmp_dir, 'day026_report.docx')\n"
            "\n"
            "reporter.to_pdf('Q1 Business Report', sections, pdf_path)\n"
            "reporter.to_docx('Q1 Business Report', sections, docx_path)\n"
            "print(f'PDF:  {pdf_path} ({os.path.getsize(pdf_path):,} bytes)')\n"
            "print(f'DOCX: {docx_path} ({os.path.getsize(docx_path):,} bytes)')"
        ),
        md("## Action 3 — Read Back and Verify"),
        code(
            "pages = read_pdf_pages(pdf_path)\n"
            "print(f'PDF: {len(pages)} page(s), first 120 chars:')\n"
            "print((pages[0] if pages else '(no text)').strip()[:120])\n"
            "\n"
            "body = read_docx_text(docx_path)\n"
            "print(f'\\nDOCX: {len(body)} chars, preview:')\n"
            "print(body[:120])\n"
            "print('\\nReport generation complete!')"
        ),
    ]


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("Generating Day 026 notebooks...")
    ex_dir = DAY_DIR / "exercises"
    proj_dir = DAY_DIR / "project"
    sol_dir = proj_dir / "solution"
    for d in (ex_dir, proj_dir, sol_dir):
        d.mkdir(parents=True, exist_ok=True)

    write_nb(ex_dir / "exercise_01.ipynb", ex01())
    write_nb(ex_dir / "exercise_02.ipynb", ex02())
    write_nb(ex_dir / "exercise_03.ipynb", ex03())
    write_nb(ex_dir / "exercise_04.ipynb", ex04())
    write_nb(ex_dir / "exercise_05.ipynb", ex05())
    write_nb(proj_dir / "project.ipynb", project_nb())
    write_nb(sol_dir / "solution.ipynb", solution_nb())
    print("Done.")


if __name__ == "__main__":
    main()
